import os
import docx

class Doc():

    def create_doc(self, location):

        doc = docx.Document()

        doc.add_heading('Project No 1', 0)

        p1 = doc.add_paragraph('Creation date: ')
        p1.add_run('14.11.2018').bold = True

        p2 = doc.add_paragraph('Type of service: ')
        p2.add_run('Translation').bold = True

        p3 = doc.add_paragraph('Project Manager: ')
        p3.add_run('Sebastjan Leskovar').bold = True

        # TODO: Tukaj dodaj prazno vrstico.

        # TODO: Change the location where the document is saved. It should be saved at root of the project folder.
        doc.save(os.path.join(location, 'Project No 1 - Log.docx'))
