import os
import docx
import datetime
from src import log_tracking, project_information

class Log_document():
    '''This class is used to create a log document and insert current time into that document.'''

    def create_log_document(self):
        '''Create a document with the project's information in the project folder. The document can be used to hold notes for this project.'''

        log = log_tracking.Logging()
        pr_info = project_information.project_info()
        pr_info.get_project_info()
        
        doc = docx.Document()

        doc.add_heading(pr_info.project_folder_name, 0)

        self.current_time()
        p1 = doc.add_paragraph('Creation date: ')
        p1.add_run(self.formatted_current_time).bold = True

        p2 = doc.add_paragraph('')
        p3 = doc.add_paragraph('')

        doc.save(pr_info.log_name_and_location)
        
        log.write_to_log("The log document was created in project folder '%s'." % pr_info.log_name)
        
        print("A log '%s' was created in project folder." % pr_info.log_name)

    def current_time(self):
        current_time = datetime.datetime.now()
        self.formatted_current_time = current_time.strftime("%d. %m. %Y, %H:%M:%S")